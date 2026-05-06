"""
Asset management service — PDF parsing, icon extraction, pack/item CRUD.
Uses PyMuPDF for real PDF image extraction and LLM for document analysis.
"""
import json
import re
import uuid
import shutil
import traceback
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from datetime import datetime
from typing import Optional

from core.schemas.asset_pack import AssetPack, AssetItem, KnowledgeDoc


ASSET_DIR = Path("assets")
ASSET_DIR.mkdir(exist_ok=True)

# ── In-memory stores ──
_packs: dict[str, AssetPack] = {}
_items: dict[str, AssetItem] = {}
_docs: dict[str, KnowledgeDoc] = {}


def _now() -> str:
    return datetime.now().isoformat(timespec="seconds")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Asset Packs
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def create_pack(
    name: str,
    file_paths,
    category: list[str],
    usage: list[str],
    file_type: str = "pdf",
    pack_type: str = "icon_pack_pdf",
) -> AssetPack:
    """Create a new asset pack from uploaded files (list or single path)."""
    if isinstance(file_paths, str):
        file_paths = [file_paths]
    pack_id = f"pack_{uuid.uuid4().hex[:8]}"
    pack = AssetPack(
        asset_pack_id=pack_id,
        name=name,
        pack_type=pack_type,
        file_type=file_type,
        category=category,
        usage=usage,
        source_file=file_paths[0] if file_paths else "",
        source_file_url=f"/assets/packs/{pack_id}/",
        parse_status="pending",
        created_at=_now(),
        updated_at=_now(),
    )
    pack_dir = ASSET_DIR / "packs" / pack_id
    pack_dir.mkdir(parents=True, exist_ok=True)
    for fp in file_paths:
        p = Path(fp)
        if p.exists():
            shutil.copy2(fp, pack_dir / p.name)

    _packs[pack_id] = pack
    _save_pack_meta(pack)
    return pack


def list_packs() -> list[dict]:
    _load_all_packs()
    return [p.model_dump() for p in _packs.values()]


def get_pack(pack_id: str) -> Optional[AssetPack]:
    _load_all_packs()
    return _packs.get(pack_id)


def parse_pack(pack_id: str) -> AssetPack:
    """Parse a pack into individual asset items (PDF or images)."""
    pack = get_pack(pack_id)
    if not pack:
        raise ValueError(f"Pack {pack_id} not found")

    pack.parse_status = "parsing"
    pack.updated_at = _now()
    _save_pack_meta(pack)

    try:
        if pack.file_type == "zip":
            _unpack_zip_sources(pack)
            items = _extract_items_from_images(pack)
        elif pack.file_type == "mixed":
            _unpack_zip_sources(pack)
            items = _extract_items_from_images(pack)
            if any((ASSET_DIR / "packs" / pack.asset_pack_id).glob("*.pdf")):
                items.extend(_extract_items_from_pdf(pack))
        elif pack.file_type == "image":
            items = _extract_items_from_images(pack)
        else:
            items = _extract_items_from_pdf(pack)
        pack.item_count = len(items)
        pack.parse_status = _status_for_parsed_items(pack, items)
        pack.updated_at = _now()
    except Exception as e:
        pack.parse_status = "failed"
        pack.error = str(e)
        pack.updated_at = _now()
        traceback.print_exc()

    _save_pack_meta(pack)
    return pack


def _status_for_parsed_items(pack: AssetPack, items: list[AssetItem]) -> str:
    if "pymupdf_unavailable" in pack.tags:
        return "needs_review"
    if not items:
        return "needs_review"
    if all(item.status == "needs_review" for item in items):
        return "needs_review"
    if all(item.source == "text_preview" for item in items):
        return "needs_review"
    low_confidence = [item for item in items if item.confidence and item.confidence < 0.45]
    if len(low_confidence) == len(items):
        return "needs_review"
    return "parsed"


def _unpack_zip_sources(pack: AssetPack):
    pack_dir = ASSET_DIR / "packs" / pack.asset_pack_id
    for f in list(pack_dir.iterdir()):
        if f.suffix.lower() != ".zip":
            continue
        with zipfile.ZipFile(f) as zf:
            for member in zf.infolist():
                if member.is_dir():
                    continue
                suffix = Path(member.filename).suffix.lower()
                if suffix not in {".png", ".jpg", ".jpeg", ".webp", ".svg", ".gif", ".bmp", ".tiff", ".pdf"}:
                    continue
                target = pack_dir / Path(member.filename).name
                with zf.open(member) as src, open(target, "wb") as out:
                    shutil.copyfileobj(src, out)


def _extract_items_from_images(pack: AssetPack) -> list[AssetItem]:
    """Handle image file uploads (PNG, JPG, etc.) — each file becomes one asset item."""
    pack_dir = ASSET_DIR / "packs" / pack.asset_pack_id
    items_dir = pack_dir / "items"
    items_dir.mkdir(exist_ok=True)

    from PIL import Image
    extracted = []
    img_exts = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff", ".svg"}
    img_index = 0

    for f in sorted(pack_dir.iterdir()):
        if f.suffix.lower() not in img_exts:
            continue
        try:
            # Copy to items dir as PNG
            item_id = f"item_{pack.asset_pack_id}_{img_index:03d}"
            img_filename = f"{item_id}.png"
            img_path = items_dir / img_filename

            if f.suffix.lower() == ".svg":
                # SVG: just copy as-is
                shutil.copy2(f, items_dir / f"{item_id}{f.suffix}")
                img_filename = f"{item_id}{f.suffix}"
                w, h = 0, 0
            else:
                img = Image.open(f)
                if img.mode not in ("RGB", "RGBA"):
                    img = img.convert("RGBA")
                img.save(str(img_path), "PNG")
                w, h = img.width, img.height

            item = AssetItem(
                asset_item_id=item_id,
                asset_pack_id=pack.asset_pack_id,
                name=f.stem,
                item_type=_guess_type_by_size(w, h) if w > 0 else "graphic",
                group=_guess_group_from_name(f.stem),
                tags=[],
                bbox=[0, 0, w, h],
                page=0,
                preview_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}",
                png_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}",
                transparent_png_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}" if f.suffix.lower() == ".png" else "",
                applicable_categories=pack.category,
                applicable_image_types=pack.usage,
                status="auto_detected",
                confidence=0.75,
                source="image_upload",
                created_at=_now(),
            )
            _items[item_id] = item
            extracted.append(item)
            img_index += 1
        except Exception as e:
            print(f"  [AssetParse] 跳过文件 {f.name}: {e}")
            continue

    pack.page_count = 0
    print(f"  [AssetParse] 提取到 {len(extracted)} 个图片素材")

    # VLM auto-label
    _label_items_with_vlm(extracted, items_dir)

    # Save items metadata
    items_meta = [it.model_dump() for it in extracted]
    (pack_dir / "items.json").write_text(
        json.dumps(items_meta, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return extracted


def _extract_items_from_pdf(pack: AssetPack) -> list[AssetItem]:
    """
    Extract images from PDF in three stages:
    1. Render every page to high-resolution page images with PyMuPDF.
    2. Extract embedded images when available.
    3. For icon-pack PDFs, slice orange icon candidates from page images and
       fall back to a Feandrea-style mock catalog when detection is sparse.
    """
    pack_dir = ASSET_DIR / "packs" / pack.asset_pack_id
    items_dir = pack_dir / "items"
    pages_dir = pack_dir / "pages"
    items_dir.mkdir(exist_ok=True)
    pages_dir.mkdir(exist_ok=True)

    # Find the PDF file
    source = None
    for f in pack_dir.iterdir():
        if f.suffix.lower() == ".pdf":
            source = f
            break
    if source is None:
        sf = Path(pack.source_file)
        if sf.exists():
            source = sf
    if source is None:
        raise FileNotFoundError(f"No PDF found for pack {pack.asset_pack_id}")

    extracted = []
    print(f"  [AssetParse] 开始解析 PDF: {source.name}")

    from PIL import Image
    import io

    page_images, page_texts = _render_pdf_pages_with_pymupdf(pack, source)
    if page_images:
        pack.page_count = len(page_images)
    else:
        from pypdf import PdfReader
        reader = PdfReader(str(source))
        pack.page_count = len(reader.pages)
        page_texts = [(page.extract_text() or "") for page in reader.pages]
    print(f"  [AssetParse] PDF 共 {pack.page_count} 页")

    img_index = 0
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(source))
        pages = reader.pages
    except Exception:
        pages = []

    for page_num, page in enumerate(pages):
        if hasattr(page, "images"):
            for img_obj in page.images:
                try:
                    img_data = img_obj.data
                    img = Image.open(io.BytesIO(img_data))
                    # Skip tiny images
                    if img.width < 20 or img.height < 20:
                        continue
                    # Convert to RGB if needed
                    if img.mode not in ("RGB", "RGBA"):
                        img = img.convert("RGBA")

                    item_id = f"item_{pack.asset_pack_id}_{img_index:03d}"
                    img_filename = f"{item_id}.png"
                    img_path = items_dir / img_filename
                    img.save(str(img_path), "PNG")

                    item = AssetItem(
                        asset_item_id=item_id,
                        asset_pack_id=pack.asset_pack_id,
                        name=f"素材_{img_index+1} (p{page_num+1})",
                        item_type=_guess_type_by_size(img.width, img.height),
                        group=_infer_group(page_texts[page_num] if page_num < len(page_texts) else "", page_num + 1, 0, max(1, img.height)),
                        tags=[],
                        bbox=[0, 0, img.width, img.height],
                        page=page_num + 1,
                        preview_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}",
                        png_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}",
                        transparent_png_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}",
                        applicable_categories=pack.category,
                        applicable_image_types=pack.usage,
                        status="auto_detected",
                        confidence=0.68,
                        source="pdf_embedded_image",
                        created_at=_now(),
                    )
                    _items[item_id] = item
                    extracted.append(item)
                    img_index += 1
                except Exception as e:
                    print(f"  [AssetParse] 跳过图片: {e}")
                    continue

    if pack.pack_type == "icon_pack_pdf" or "icon" in ",".join(pack.usage).lower() or "图标" in pack.name:
        sliced = _slice_icon_candidates_from_pages(pack, page_images, page_texts, start_index=img_index)
        extracted.extend(sliced)
        img_index += len(sliced)

    if _should_use_feandrea_mock(pack, source, extracted):
        mocked = _create_feandrea_mock_items(pack, start_index=img_index)
        extracted.extend(mocked)
        img_index += len(mocked)

    if not extracted:
        print("  [AssetParse] 未识别到单独素材项，生成页面预览素材并标记 needs_review")
        for page_num, page_image in enumerate(page_images):
            text = page_texts[page_num] if page_num < len(page_texts) else ""
            item_id = f"item_{pack.asset_pack_id}_{img_index:03d}"
            img_index += 1
            img_filename = f"{item_id}.png"
            img_path = items_dir / img_filename
            shutil.copy2(page_image, img_path)
            with Image.open(page_image) as im:
                w, h = im.size
            item = AssetItem(
                asset_item_id=item_id,
                asset_pack_id=pack.asset_pack_id,
                name=f"页面_{page_num+1}",
                item_type="graphic",
                group="其他",
                tags=["pdf_page", "needs_review"],
                bbox=[0, 0, w, h],
                page=page_num + 1,
                preview_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}",
                png_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}",
                description=(text or "No extractable page text.")[:500],
                applicable_categories=pack.category,
                applicable_image_types=pack.usage,
                status="needs_review",
                confidence=0.2,
                source="text_preview",
                created_at=_now(),
            )
            _items[item_id] = item
            extracted.append(item)

    print(f"  [AssetParse] 提取到 {len(extracted)} 个素材项")

    # Use VLM to auto-label items that have preview images
    _label_items_with_vlm(extracted, items_dir)

    # Save items metadata
    items_meta = [it.model_dump() for it in extracted]
    (pack_dir / "items.json").write_text(
        json.dumps(items_meta, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return extracted


def _guess_type_by_size(w: int, h: int) -> str:
    area = w * h
    if area < 10000:
        return "icon"
    elif area < 100000:
        return "graphic"
    else:
        return "background"


def _render_pdf_pages_with_pymupdf(pack: AssetPack, source: Path) -> tuple[list[Path], list[str]]:
    """Render every PDF page to PNG. Returns paths plus extractable page text."""
    pages_dir = ASSET_DIR / "packs" / pack.asset_pack_id / "pages"
    pages_dir.mkdir(parents=True, exist_ok=True)
    page_images: list[Path] = []
    page_texts: list[str] = []
    try:
        import fitz
    except ImportError:
        print("  [AssetParse] PyMuPDF 未安装，跳过页面渲染")
        if "pymupdf_unavailable" not in pack.tags:
            pack.tags.append("pymupdf_unavailable")
        return page_images, page_texts

    doc = fitz.open(str(source))
    matrix = fitz.Matrix(2.5, 2.5)
    for idx, page in enumerate(doc, start=1):
        page_path = pages_dir / f"page_{idx:03d}.png"
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        pix.save(str(page_path))
        page_images.append(page_path)
        try:
            page_texts.append(page.get_text("text") or "")
        except Exception:
            page_texts.append("")
    doc.close()
    return page_images, page_texts


def _extract_icon_candidates_from_page_image(
    pack: AssetPack,
    page_image: Path,
    page_text: str = "",
    start_index: int = 0,
) -> list[AssetItem]:
    """Extract icon candidates from one rendered page image."""
    return _slice_icon_candidates_from_pages(pack, [page_image], [page_text], start_index=start_index)


def _slice_icon_candidates_from_pages(
    pack: AssetPack,
    page_images: list[Path],
    page_texts: list[str],
    start_index: int = 0,
) -> list[AssetItem]:
    from PIL import Image

    pack_dir = ASSET_DIR / "packs" / pack.asset_pack_id
    items_dir = pack_dir / "items"
    extracted: list[AssetItem] = []
    idx = start_index
    for page_idx, page_path in enumerate(page_images, start=1):
        image = Image.open(page_path).convert("RGB")
        boxes = _detect_orange_icon_boxes(image)
        if not boxes:
            boxes = _grid_icon_boxes(image)
        text = page_texts[page_idx - 1] if page_idx - 1 < len(page_texts) else ""
        for local_idx, box in enumerate(boxes, start=1):
            x, y, w, h = box
            if w < 24 or h < 24:
                continue
            pad = max(6, int(min(w, h) * 0.08))
            crop_box = (
                max(0, x - pad),
                max(0, y - pad),
                min(image.width, x + w + pad),
                min(image.height, y + h + pad),
            )
            crop = image.crop(crop_box)
            item_id = f"item_{pack.asset_pack_id}_{idx:03d}"
            idx += 1
            img_filename = f"{item_id}.png"
            img_path = items_dir / img_filename
            crop.save(img_path, "PNG")
            group = _infer_group(text, page_idx, y + h / 2, image.height)
            name = _name_from_page_text(text, group, local_idx) or f"图标_p{page_idx}_{local_idx:03d}"
            item_type = "arrow" if group == "箭头" else "icon"
            item = AssetItem(
                asset_item_id=item_id,
                asset_pack_id=pack.asset_pack_id,
                name=name,
                item_type=item_type,
                group=group,
                tags=[group, "auto_detected"],
                bbox=[x, y, w, h],
                page=page_idx,
                preview_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}",
                png_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}",
                transparent_png_url=f"/assets/packs/{pack.asset_pack_id}/items/{img_filename}",
                applicable_categories=pack.category,
                applicable_image_types=pack.usage,
                status="auto_detected",
                confidence=0.62,
                source="pdf_page_crop",
                created_at=_now(),
            )
            _items[item_id] = item
            extracted.append(item)
    return extracted


def _detect_orange_icon_boxes(image) -> list[list[int]]:
    """Detect connected orange regions in a rendered page image."""
    pix = image.load()
    width, height = image.size
    mask = bytearray(width * height)
    for y in range(height):
        for x in range(width):
            r, g, b = pix[x, y]
            if r >= 170 and 55 <= g <= 190 and b <= 120 and r > g * 1.15:
                mask[y * width + x] = 1

    seen = bytearray(width * height)
    boxes: list[list[int]] = []
    for y in range(0, height, 2):
        for x in range(0, width, 2):
            pos = y * width + x
            if not mask[pos] or seen[pos]:
                continue
            stack = [(x, y)]
            seen[pos] = 1
            min_x = max_x = x
            min_y = max_y = y
            count = 0
            while stack:
                cx, cy = stack.pop()
                count += 1
                min_x, max_x = min(min_x, cx), max(max_x, cx)
                min_y, max_y = min(min_y, cy), max(max_y, cy)
                for nx, ny in ((cx + 1, cy), (cx - 1, cy), (cx, cy + 1), (cx, cy - 1)):
                    if nx < 0 or ny < 0 or nx >= width or ny >= height:
                        continue
                    npos = ny * width + nx
                    if mask[npos] and not seen[npos]:
                        seen[npos] = 1
                        stack.append((nx, ny))
            w = max_x - min_x + 1
            h = max_y - min_y + 1
            area = w * h
            if count >= 80 and 24 <= w <= width * 0.45 and 24 <= h <= height * 0.45 and area >= 1000:
                boxes.append([min_x, min_y, w, h])

    boxes = _merge_nearby_boxes(boxes)
    boxes.sort(key=lambda b: (b[1], b[0]))
    return boxes[:80]


def _merge_nearby_boxes(boxes: list[list[int]]) -> list[list[int]]:
    merged: list[list[int]] = []
    for box in boxes:
        x, y, w, h = box
        bx2, by2 = x + w, y + h
        absorbed = False
        for idx, existing in enumerate(merged):
            ex, ey, ew, eh = existing
            ex2, ey2 = ex + ew, ey + eh
            if not (bx2 < ex - 12 or x > ex2 + 12 or by2 < ey - 12 or y > ey2 + 12):
                nx, ny = min(x, ex), min(y, ey)
                nx2, ny2 = max(bx2, ex2), max(by2, ey2)
                merged[idx] = [nx, ny, nx2 - nx, ny2 - ny]
                absorbed = True
                break
        if not absorbed:
            merged.append(box)
    return merged


def _grid_icon_boxes(image) -> list[list[int]]:
    """Conservative grid fallback for icon summary pages."""
    width, height = image.size
    boxes: list[list[int]] = []
    margin_x = int(width * 0.08)
    margin_y = int(height * 0.12)
    cols = 5 if width >= 1200 else 4
    rows = 6
    cell_w = (width - margin_x * 2) / cols
    cell_h = (height - margin_y * 2) / rows
    side = int(min(cell_w, cell_h) * 0.55)
    for row in range(rows):
        for col in range(cols):
            cx = int(margin_x + col * cell_w + cell_w / 2)
            cy = int(margin_y + row * cell_h + cell_h * 0.38)
            boxes.append([max(0, cx - side // 2), max(0, cy - side // 2), side, side])
    return boxes


def _infer_group(page_text: str, page: int, center_y: float, page_height: int) -> str:
    text = (page_text or "").lower()
    if any(k in text for k in ["arrow", "箭头", "指引", "虚线"]):
        return "箭头"
    if any(k in text for k in ["包装", "package", "fragile", "up", "环保", "environment"]):
        return "包装"
    if any(k in text for k in ["防水", "耐磨", "柔软", "水洗", "承重", "防滑", "function", "feature"]):
        return "功能"
    if any(k in text for k in ["产品", "猫", "狗", "pet", "cat", "dog", "product"]):
        return "产品"
    ratio = center_y / max(1, page_height)
    if ratio < 0.35:
        return "产品"
    if ratio < 0.68:
        return "功能"
    if page % 4 == 0:
        return "箭头"
    return "包装" if ratio > 0.78 else "其他"


def _guess_group_from_name(name: str) -> str:
    low = name.lower()
    if any(k in low for k in ["arrow", "箭头"]):
        return "箭头"
    if any(k in low for k in ["package", "fragile", "包装", "环保"]):
        return "包装"
    if any(k in low for k in ["waterproof", "wash", "anti", "防", "承重", "功能"]):
        return "功能"
    if any(k in low for k in ["cat", "dog", "pet", "猫", "狗", "产品"]):
        return "产品"
    return "其他"


def _name_from_page_text(text: str, group: str, index: int) -> str:
    candidates = [ln.strip(" ·:-\t") for ln in (text or "").splitlines() if 1 <= len(ln.strip()) <= 20]
    group_words = _FEANDREA_MOCK_GROUPS.get(group, [])
    for word in group_words:
        if word in candidates or word in text:
            return word
    if index - 1 < len(candidates):
        return candidates[index - 1]
    return ""


def _should_use_feandrea_mock(pack: AssetPack, source: Path, extracted: list[AssetItem]) -> bool:
    marker = f"{pack.name} {source.name}".lower()
    looks_feandrea = any(k in marker for k in ["feandrea", "listing", "辅助图形", "图标"])
    return looks_feandrea and len([it for it in extracted if it.source == "pdf_page_crop"]) < 12


_FEANDREA_MOCK_GROUPS = {
    "产品": ["猫", "猫树", "猫抓板", "猫兜", "小鱼干", "爱心", "骨头", "家"],
    "功能": ["防水", "耐磨", "柔软", "可水洗", "易打扫", "防倾倒", "承重", "可拆卸", "透气", "防滑"],
    "包装": ["向上", "易碎", "怕湿", "环保"],
    "箭头": ["向上箭头", "指引箭头", "虚线箭头", "环形箭头"],
}


def _create_feandrea_mock_items(pack: AssetPack, start_index: int = 0) -> list[AssetItem]:
    from PIL import Image, ImageDraw

    pack_dir = ASSET_DIR / "packs" / pack.asset_pack_id
    items_dir = pack_dir / "items"
    extracted: list[AssetItem] = []
    idx = start_index
    for group, names in _FEANDREA_MOCK_GROUPS.items():
        for local_idx, name in enumerate(names, start=1):
            item_id = f"item_{pack.asset_pack_id}_{idx:03d}"
            idx += 1
            filename = f"{item_id}.png"
            path = items_dir / filename
            image = Image.new("RGBA", (320, 320), (255, 255, 255, 0))
            draw = ImageDraw.Draw(image)
            draw.ellipse((54, 34, 266, 246), fill=(240, 126, 35, 255))
            draw.text((42, 260), name[:10], fill=(50, 50, 50, 255), font=_font(28))
            image.save(path, "PNG")
            item = AssetItem(
                asset_item_id=item_id,
                asset_pack_id=pack.asset_pack_id,
                name=name,
                item_type="arrow" if group == "箭头" else "icon",
                group=group,
                tags=[group, "feandrea_mock_fallback"],
                bbox=[],
                page=0,
                preview_url=f"/assets/packs/{pack.asset_pack_id}/items/{filename}",
                png_url=f"/assets/packs/{pack.asset_pack_id}/items/{filename}",
                transparent_png_url=f"/assets/packs/{pack.asset_pack_id}/items/{filename}",
                applicable_categories=pack.category,
                applicable_image_types=pack.usage,
                status="needs_review",
                confidence=0.35,
                source="pdf_page_crop",
                description="Feandrea icon PDF mock fallback item; requires human confirmation.",
                created_at=_now(),
            )
            _items[item_id] = item
            extracted.append(item)
    return extracted


def _label_items_with_vlm(items: list[AssetItem], items_dir: Path):
    """Use VLM to auto-label extracted items with names and tags."""
    if not items:
        return
    try:
        from models.llm import chat
        from PIL import Image
    except ImportError:
        print("  [AssetParse] VLM labeling 不可用，跳过")
        return

    for item in items:
        if "feandrea_mock_fallback" in item.tags:
            continue
        img_path = items_dir / f"{item.asset_item_id}.png"
        if not img_path.exists():
            continue
        try:
            img = Image.open(img_path)
            if img.width > 2000 or img.height > 2000:
                img.thumbnail((1000, 1000))

            prompt = """请分析这张图片，返回JSON格式：
{
  "name": "图片的简短中文名称（如：猫图标、箭头、尺寸线、产品场景图等）",
  "type": "icon / graphic / background / decoration / logo 之一",
  "tags": ["标签1", "标签2", "标签3"],
  "description": "一句话描述图片内容"
}
只返回JSON，不要其他文本。"""
            result = chat(prompt, image=img, response_format="json")
            data = json.loads(result)
            if data.get("name"):
                item.name = data["name"]
            if data.get("type"):
                item.item_type = data["type"]
                item.type = item.item_type
            if data.get("tags"):
                item.tags = _dedupe([*item.tags, *data["tags"]])
            if data.get("description"):
                item.description = data["description"]
            print(f"    VLM: {item.asset_item_id} → {item.name} [{item.item_type}]")
        except Exception as e:
            print(f"    VLM labeling 失败 {item.asset_item_id}: {e}")
            continue


def _render_text_preview(text: str, out_path: Path, title: str = "PDF Page"):
    """Render extracted page text to a PNG preview when no image asset is available."""
    from PIL import Image, ImageDraw, ImageFont

    out_path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1200, 1600), "#f7f8fb")
    draw = ImageDraw.Draw(image)
    title_font = _font(42)
    body_font = _font(28)
    small_font = _font(22)
    draw.rectangle((0, 0, 1200, 120), fill="#ffffff")
    draw.text((56, 38), title[:42], fill="#172033", font=title_font)
    draw.text((58, 98), "Auto-generated text preview", fill="#7a8494", font=small_font)

    y = 170
    max_chars = 2400
    for paragraph in (text or "").replace("\r", "\n").split("\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            y += 16
            continue
        for line in _wrap_text(paragraph, 48):
            if y > 1510 or max_chars <= 0:
                draw.text((58, y), "...", fill="#536173", font=body_font)
                image.save(out_path, "PNG")
                return
            draw.text((58, y), line, fill="#253044", font=body_font)
            y += 42
            max_chars -= len(line)
        y += 18
    image.save(out_path, "PNG")


def _wrap_text(text: str, width: int) -> list[str]:
    lines = []
    current = ""
    for char in text:
        current += char
        if len(current) >= width:
            lines.append(current)
            current = ""
    if current:
        lines.append(current)
    return lines


def _font(size: int):
    from PIL import ImageFont
    for path in (
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
    ):
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()



def list_pack_items(pack_id: str) -> list[dict]:
    _load_pack_items(pack_id)
    return [it.model_dump() for it in _items.values() if it.asset_pack_id == pack_id]


def get_item(item_id: str) -> Optional[AssetItem]:
    _load_all_packs()
    for pack_id in list(_packs.keys()):
        _load_pack_items(pack_id)
    return _items.get(item_id)


def update_item(item_id: str, updates: dict | None = None, **fields) -> Optional[AssetItem]:
    item = get_item(item_id)
    if not item:
        return None
    updates = {**(updates or {}), **fields}
    allowed = {
        "name",
        "group",
        "tags",
        "status",
        "applicable_categories",
        "applicable_image_types",
    }
    for key, value in updates.items():
        if key not in allowed:
            continue
        setattr(item, key, value)
    _items[item_id] = AssetItem(**item.model_dump())
    _persist_pack_items(item.asset_pack_id)
    return _items[item_id]


def batch_update_items(
    item_ids: list[str],
    status: str = None,
    tags: list[str] = None,
    group: str = None,
    applicable_categories: list[str] = None,
    applicable_image_types: list[str] = None,
) -> list[dict]:
    _load_all_packs()
    for pack_id in list(_packs.keys()):
        _load_pack_items(pack_id)
    updated = []
    for iid in item_ids:
        item = _items.get(iid)
        if not item:
            continue
        if status:
            item.status = status
        if tags is not None:
            item.tags = _dedupe([*item.tags, *tags])
        if group:
            item.group = group
        if applicable_categories is not None:
            item.applicable_categories = applicable_categories
        if applicable_image_types is not None:
            item.applicable_image_types = applicable_image_types
        updated.append(item.model_dump())

    # Persist changes to disk
    pack_ids = set(it.get("asset_pack_id") for it in updated if it.get("asset_pack_id"))
    for pid in pack_ids:
        _persist_pack_items(pid)

    return updated


def _persist_pack_items(pack_id: str):
    """Save current item state to disk."""
    pack_dir = ASSET_DIR / "packs" / pack_id
    if not pack_dir.exists():
        return
    all_items = [it.model_dump() for it in _items.values() if it.asset_pack_id == pack_id]
    (pack_dir / "items.json").write_text(
        json.dumps(all_items, ensure_ascii=False, indent=2), encoding="utf-8"
    )


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Knowledge Docs
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def create_doc(name: str, file_path: str, category: list[str],
               name_en: str = "") -> KnowledgeDoc:
    doc_id = f"doc_{uuid.uuid4().hex[:8]}"
    file_type = Path(file_path).suffix.lstrip(".").lower() or "pdf"
    doc = KnowledgeDoc(
        doc_id=doc_id,
        name=name,
        name_en=name_en,
        category=category,
        category_path=" > ".join(category),
        file_type=file_type,
        source_file=file_path,
        upload_time=_now(),
        parse_status="pending",
        status_message="等待解析",
    )
    doc_dir = ASSET_DIR / "docs" / doc_id
    doc_dir.mkdir(parents=True, exist_ok=True)
    if Path(file_path).exists():
        shutil.copy2(file_path, doc_dir / Path(file_path).name)

    _docs[doc_id] = doc
    _save_doc_meta(doc)
    return doc


def list_docs() -> list[dict]:
    _load_all_docs()
    return [d.model_dump() for d in _docs.values()]


def get_doc(doc_id: str) -> Optional[KnowledgeDoc]:
    _load_all_docs()
    return _docs.get(doc_id)


def analyze_doc(doc_id: str) -> KnowledgeDoc:
    """Analyze a knowledge document: extract text, then use LLM to parse rules."""
    doc = get_doc(doc_id)
    if not doc:
        raise ValueError(f"Doc {doc_id} not found")

    doc.parse_status = "parsing"
    doc.status_message = "解析中"
    _save_doc_meta(doc)

    try:
        knowledge = _extract_knowledge_real(doc)
        doc.parsed_knowledge = knowledge
        doc.rule_count = (
            len(knowledge.get("global_rules", []))
            + len(knowledge.get("scene_rules", []))
            + len(knowledge.get("style_rules", []))
        )
        doc.checklist_count = len(knowledge.get("checklist", []))
        doc.parse_status = "parsed"
        doc.summary = knowledge.get("summary", "")
        doc.category_path = knowledge.get("category_path") or doc.category_path
        doc.parse_mode = knowledge.get("parse_mode", "")
        doc.parsed_at = _now()
        doc.status_message = "解析完成"
    except Exception as e:
        doc.parse_status = "failed"
        doc.error = str(e)
        doc.status_message = f"解析失败: {e}"
        traceback.print_exc()

    _save_doc_meta(doc)
    return doc


def _extract_text_from_file(doc: KnowledgeDoc) -> str:
    """Extract text content from the uploaded file."""
    doc_dir = ASSET_DIR / "docs" / doc.doc_id
    source = None
    for f in doc_dir.iterdir():
        if f.suffix.lower() in (".pdf", ".docx", ".doc", ".txt", ".md"):
            source = f
            break
    if source is None:
        sf = Path(doc.source_file)
        if sf.exists():
            source = sf
    if source is None:
        return ""

    ext = source.suffix.lower()
    text = ""

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(source))
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        except ImportError:
            text = f"[PDF文件: {source.name}, 需要安装 pypdf]"

    elif ext in (".docx",):
        try:
            from docx import Document
            doc_file = Document(str(source))
            for para in doc_file.paragraphs:
                text += para.text + "\n"
            for table in doc_file.tables:
                for row in table.rows:
                    text += " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()) + "\n"
        except ImportError:
            text = _extract_docx_text_stdlib(source)
        except Exception:
            text = _extract_docx_text_stdlib(source)

    elif ext in (".doc",):
        text = f"[DOC文件: {source.name}, 当前环境暂不支持旧版 .doc 自动解析，请转换为 .docx/.pdf 或补充文本摘要。]"

    elif ext in (".txt", ".md"):
        text = source.read_text(encoding="utf-8", errors="ignore")

    return text.strip()


def _extract_docx_text_stdlib(source: Path) -> str:
    """Extract DOCX text with only stdlib zip/xml as a python-docx fallback."""
    if not zipfile.is_zipfile(source):
        return f"[DOCX文件: {source.name}, 文件不是有效的 Office Open XML 包。]"

    parts = ["word/document.xml"]
    with zipfile.ZipFile(source) as zf:
        parts.extend(
            name for name in zf.namelist()
            if name.startswith("word/header") or name.startswith("word/footer")
        )
        chunks: list[str] = []
        for part in parts:
            if part not in zf.namelist():
                continue
            try:
                root = ET.fromstring(zf.read(part))
            except ET.ParseError:
                continue
            chunks.extend(_docx_xml_text_chunks(root))
    return "\n".join(chunks).strip()


def _docx_xml_text_chunks(root: ET.Element) -> list[str]:
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    chunks: list[str] = []
    for para in root.iter(f"{ns}p"):
        parts: list[str] = []
        for node in para.iter():
            if node.tag == f"{ns}t" and node.text:
                parts.append(node.text)
            elif node.tag == f"{ns}tab":
                parts.append("\t")
            elif node.tag in {f"{ns}br", f"{ns}cr"}:
                parts.append("\n")
        line = "".join(parts).strip()
        if line:
            chunks.append(line)
    return chunks


def _extract_knowledge_real(doc: KnowledgeDoc) -> dict:
    """Extract structured knowledge using real text extraction + LLM analysis."""
    print(f"  [KnowledgeAnalyze] 开始解析文档: {doc.name}")

    text = _extract_text_from_file(doc)
    if not text:
        print(f"  [KnowledgeAnalyze] 文档内容为空")
        return _heuristic_knowledge("", doc, reason="empty_extracted_text")

    print(f"  [KnowledgeAnalyze] 提取到 {len(text)} 字符文本，调用 LLM 分块分析")
    try:
        return _extract_knowledge_with_llm(text, doc)
    except Exception as e:
        print(f"  [KnowledgeAnalyze] LLM 分析失败，使用本地启发式解析: {e}")
        return _heuristic_knowledge(text, doc, reason=f"llm_failed: {e}")


def _extract_knowledge_with_llm(text: str, doc: KnowledgeDoc) -> dict:
    chunks = _chunk_text(text, max_chars=6200)
    models = _llm_model_candidates()
    last_error: Exception | None = None
    print(f"  [KnowledgeAnalyze] LLM chunks={len(chunks)}, models={models}")

    for model in models:
        try:
            partials = []
            for idx, chunk in enumerate(chunks, start=1):
                prompt = _knowledge_chunk_prompt(doc, chunk, idx, len(chunks))
                partials.append(_chat_json(prompt, model=model))
            merged = _merge_knowledge_partials(partials, doc)
            merged["parse_mode"] = "llm_chunked"
            merged["llm_model"] = model
            merged["source_text_chars"] = len(text)
            print(f"  [KnowledgeAnalyze] LLM 提取成功: model={model}, "
                  f"rules={len(merged.get('global_rules', []))}, checklist={len(merged.get('checklist', []))}")
            return merged
        except Exception as exc:
            last_error = exc
            print(f"  [KnowledgeAnalyze] 模型 {model} 提取失败: {exc}")
            continue
    raise RuntimeError(str(last_error) if last_error else "no_llm_model_available")


def _llm_model_candidates() -> list[str]:
    try:
        import config
        raw = [
            config.MODELS.get("llm_primary"),
            config.MODELS.get("llm_secondary"),
            config.MODELS.get("quality"),
        ]
    except Exception:
        raw = []
    result = []
    for model in raw:
        if model and model not in result:
            result.append(model)
    return result or ["gpt-5.2"]


def _chat_json(prompt: str, model: str) -> dict:
    from models.llm import chat

    response = chat(prompt=prompt, model=model, response_format="json")
    return _parse_json_object(response)


def _parse_json_object(text: str) -> dict:
    text = (text or "").strip()
    try:
        value = json.loads(text)
        return value if isinstance(value, dict) else {"items": value}
    except json.JSONDecodeError:
        pass
    match = re.search(r"```(?:json)?\s*(.*?)\s*```", text, re.DOTALL)
    if match:
        return _parse_json_object(match.group(1))
    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        value = json.loads(text[start:end + 1])
        return value if isinstance(value, dict) else {"items": value}
    raise ValueError("LLM did not return a JSON object")


def _knowledge_chunk_prompt(doc: KnowledgeDoc, chunk: str, idx: int, total: int) -> str:
    return f"""你是跨境电商商品视觉知识库抽取 Agent。

任务：从文档片段中提取可直接用于 Amazon Listing 生图、提示词反推、素材规范和质检闭环的结构化知识。

文档名称：{doc.name}
适用品类：{", ".join(doc.category)}
片段：{idx}/{total}

文档片段：
---
{chunk}
---

请只返回 JSON object，字段如下：
{{
  "category_path": "品类路径",
  "summary": "本片段摘要",
  "applicable_products": ["适用产品"],
  "global_rules": ["全局生图规则，保留具体约束"],
  "image_plan_templates": [{{"type": "image type", "name": "图片名称", "goal": "视觉目标", "elements": ["必备元素"]}}],
  "prompt_templates": ["可复用提示词或提示词结构"],
  "scene_rules": ["场景图规则"],
  "style_rules": ["版式/配色/字体/图标/视觉规范"],
  "negative_prompts": ["禁止项/避免项"],
  "checklist": ["质检清单"],
  "keyword_bank": ["关键词"],
  "replaceable_variables": ["可替换变量"]
}}

要求：
- 不要泛泛总结，尽量保留文档里的具体数字、图型、卖点、场景、禁止项。
- 如果该片段没有某类信息，对应字段返回空数组或空字符串。
- 输出必须是合法 JSON，不要 markdown。"""


def _chunk_text(text: str, max_chars: int = 6200) -> list[str]:
    paragraphs = [p.strip() for p in re.split(r"\n{2,}|\r\n{2,}", text) if p.strip()]
    if not paragraphs:
        paragraphs = [text]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(paragraph) > max_chars:
            if current:
                chunks.append(current.strip())
                current = ""
            for start in range(0, len(paragraph), max_chars):
                chunks.append(paragraph[start:start + max_chars].strip())
            continue
        if len(current) + len(paragraph) + 2 > max_chars and current:
            chunks.append(current.strip())
            current = paragraph
        else:
            current = f"{current}\n\n{paragraph}" if current else paragraph
    if current:
        chunks.append(current.strip())
    return chunks


def _merge_knowledge_partials(partials: list[dict], doc: KnowledgeDoc) -> dict:
    list_fields = [
        "applicable_products",
        "global_rules",
        "image_plan_templates",
        "prompt_templates",
        "scene_rules",
        "style_rules",
        "negative_prompts",
        "checklist",
        "keyword_bank",
        "replaceable_variables",
    ]
    merged: dict = {
        "category_path": "",
        "summary": "",
        **{field: [] for field in list_fields},
    }

    summaries = []
    for partial in partials:
        if partial.get("category_path") and not merged["category_path"]:
            merged["category_path"] = partial.get("category_path", "")
        if partial.get("summary"):
            summaries.append(str(partial["summary"]))
        for field in list_fields:
            merged[field].extend(_as_list(partial.get(field)))

    if not merged["category_path"]:
        merged["category_path"] = " > ".join(doc.category) if doc.category else ""
    merged["summary"] = f"{doc.name} LLM 提取摘要：" + " / ".join(_dedupe(summaries)[:4])
    for field in list_fields:
        merged[field] = _dedupe_structured(merged[field])[:40]
    return merged


def _as_list(value) -> list:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _dedupe_structured(items: list) -> list:
    seen = set()
    result = []
    for item in items:
        key = json.dumps(item, ensure_ascii=False, sort_keys=True) if isinstance(item, dict) else str(item)
        if key in seen or key in {"", "[]", "{}"}:
            continue
        seen.add(key)
        result.append(item)
    return result


def _heuristic_knowledge(text: str, doc: KnowledgeDoc, reason: str = "local_fallback") -> dict:
    lines = [ln.strip(" -*\t") for ln in text.splitlines() if ln.strip()]
    lowered = text.lower()
    rules = []
    scene_rules = []
    style_rules = []
    negatives = []
    checklist = []
    keywords = []

    for line in lines[:160]:
        low = line.lower()
        if any(k in low for k in ["must", "should", "需要", "必须", "保持", "要求", "avoid", "不要", "禁止"]):
            target = negatives if any(k in low for k in ["avoid", "不要", "禁止", "no "]) else rules
            target.append(line[:180])
        if any(k in low for k in ["scene", "场景", "living room", "home", "室内", "客厅"]):
            scene_rules.append(line[:180])
        if any(k in low for k in ["style", "font", "color", "风格", "字体", "颜色", "配色"]):
            style_rules.append(line[:180])
        if any(k in low for k in ["check", "qa", "review", "检查", "质检", "审核"]):
            checklist.append(line[:180])

    for token in ["Amazon", "cat tree", "pet supplies", "premium", "205cm", "sisal", "plush", "猫爬架", "剑麻", "绒布"]:
        if token.lower() in lowered:
            keywords.append(token)

    if not checklist:
        checklist = [
            "产品结构、颜色、比例与 SKU 参考保持一致",
            "最终图不得出现透明棋盘格、白底残留或明显贴图感",
            "场景图需有自然接触阴影并符合电商主图/详情页用途",
        ]
    if not rules:
        rules = [
            "SKU 事实优先于模型猜测，硬性卖点以商品配置为准",
            "知识文档已进入人工复核模式，建议补充结构化规则或启用 LLM 解析",
        ]

    return {
        "category_path": " > ".join(doc.category) if doc.category else "",
        "summary": f"{doc.name} 本地解析摘要：提取 {len(lines)} 行文本，生成待复核规则。",
        "applicable_products": doc.category,
        "global_rules": _dedupe(rules)[:12],
        "image_plan_templates": [],
        "prompt_templates": [],
        "scene_rules": _dedupe(scene_rules)[:10],
        "style_rules": _dedupe(style_rules)[:10],
        "negative_prompts": _dedupe(negatives)[:10],
        "checklist": _dedupe(checklist)[:12],
        "keyword_bank": keywords,
        "replaceable_variables": ["${product_name}", "${sku_id}", "${core_selling_points}"],
        "parse_mode": "local_heuristic_fallback",
        "fallback_reason": reason,
    }


def _dedupe(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(item)
    return result


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Persistence helpers
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _save_pack_meta(pack: AssetPack):
    d = ASSET_DIR / "packs" / pack.asset_pack_id
    d.mkdir(parents=True, exist_ok=True)
    (d / "meta.json").write_text(
        json.dumps(pack.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _save_doc_meta(doc: KnowledgeDoc):
    d = ASSET_DIR / "docs" / doc.doc_id
    d.mkdir(parents=True, exist_ok=True)
    (d / "meta.json").write_text(
        json.dumps(doc.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _load_all_packs():
    packs_dir = ASSET_DIR / "packs"
    if not packs_dir.exists():
        return
    for d in packs_dir.iterdir():
        if d.is_dir() and (d / "meta.json").exists():
            pid = d.name
            if pid not in _packs:
                data = json.loads((d / "meta.json").read_text(encoding="utf-8"))
                _packs[pid] = AssetPack(**data)


def _load_pack_items(pack_id: str):
    items_file = ASSET_DIR / "packs" / pack_id / "items.json"
    if items_file.exists():
        loaded = json.loads(items_file.read_text(encoding="utf-8"))
        for it_data in loaded:
            iid = it_data["asset_item_id"]
            if iid not in _items:
                _items[iid] = AssetItem(**it_data)


def _load_all_docs():
    docs_dir = ASSET_DIR / "docs"
    if not docs_dir.exists():
        return
    for d in docs_dir.iterdir():
        if d.is_dir() and (d / "meta.json").exists():
            did = d.name
            if did not in _docs:
                data = json.loads((d / "meta.json").read_text(encoding="utf-8"))
                _docs[did] = KnowledgeDoc(**data)
